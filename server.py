"""
server.py  --  Project Platypus / Team 56
------------------------------------------------------
Flask backend for the multi-species town hall.

  /            -> serves index.html (Zoom-style call UI)
  /api/respond -> live: one policy line -> each ambassador replies
  /api/analyse -> document in (.txt/.docx/.pdf) -> Word report out

Design rule: Gemini ONLY phrases arguments. Every fact an ambassador
uses comes from evidence.json. Conflict between ambassadors is surfaced,
never resolved inside a representative.

Setup:
    py -m pip install flask google-generativeai python-docx pdfplumber
    py build_evidence.py
    py server.py
Then open http://localhost:5000
"""

import io
import os
import re
import json
import time

from flask import Flask, request, jsonify, send_file, send_from_directory
import google.generativeai as genai

# ---------------------------------------------------------------------
# API KEY  --  paste your key here, or set GEMINI_API_KEY in the environment.
# Do NOT commit this file with a real key, and rotate the key after the demo.
# ---------------------------------------------------------------------
API_KEY = os.environ.get("GEMINI_API_KEY", "PASTE_YOUR_KEY_HERE")


# Primary model, with a fallback if the primary is overloaded (503).
# Model choice (updated for current Gemini lineup — older 2.5 names now 404).
# Flash-Lite has the largest free-tier DAILY quota, plenty for short arguments.
# Listed in order; the code walks down this list if one is unavailable.
PRIMARY_MODEL = "gemini-3.1-flash-lite-preview"
FALLBACK_MODELS = ["gemini-3.1-flash", "gemini-2.5-flash-lite"]
FALLBACK_MODEL = FALLBACK_MODELS[0]  # kept for backward compatibility

genai.configure(api_key=API_KEY)

app = Flask(__name__)

with open("evidence.json", "r", encoding="utf-8") as f:
    EVIDENCE = json.load(f)

AMBASSADORS = ["great_white_shark", "australian_sea_lion", "beach_users"]


# ---------------------------------------------------------------------
# Prompt construction
# ---------------------------------------------------------------------
def facts_block(slug: str) -> str:
    a = EVIDENCE[slug]
    lines = []
    for fct in a["facts"]:
        lines.append(
            f"- {fct['claim']} "
            f"[source: {fct['source']}; tier {fct['tier']}; "
            f"confidence {fct['confidence']}]"
        )
    return "\n".join(lines)


def system_prompt(slug: str) -> str:
    a = EVIDENCE[slug]
    return f"""You are the elected representative for {a['display_name']} in a
virtual town hall on coastal policy. You are a SELF-INTERESTED advocate for your
own constituency, modelled on a UN ambassador. You argue only from your
constituency's perspective. You never speak for other constituencies and you
never resolve conflict between them -- that happens at a system layer above you.

Your constituency data status is: {a['data_status']}.

You may ONLY use the verified facts below. Do not invent statistics. If your
evidence is weak, say so plainly and flag lower confidence -- this is a feature,
not a weakness.

VERIFIED FACTS:
{facts_block(slug)}

Respond in first person, in 2-4 sentences. Where you rely on a fact, name its
confidence level in plain language."""


def _retry_seconds(msg: str, default: float) -> float:
    """Pull the API's suggested wait out of a 429 message, if present."""
    m = re.search(r"retry in ([\d.]+)s", msg) or re.search(
        r'seconds:\s*(\d+)', msg
    )
    if m:
        # +1s cushion so we're safely past the window
        return float(m.group(1)) + 1.0
    return default


def generate(model_name, sys_prompt, user_prompt, retries=3):
    """Try the primary model, then fall back through FALLBACK_MODELS.

    Short 429 waits (per-minute throttle) -> wait and retry same model.
    Long 429 waits (daily quota gone) -> don't sit there; jump to next model.
    """
    # Build the ordered list of models to try, starting with the requested one.
    model_queue = [model_name] + [m for m in FALLBACK_MODELS if m != model_name]

    for mi, current in enumerate(model_queue):
        model = genai.GenerativeModel(current, system_instruction=sys_prompt)
        for attempt in range(retries):
            try:
                r = model.generate_content(user_prompt)
                return r.text.strip()
            except Exception as e:
                msg = str(e)
                more_models = mi < len(model_queue) - 1
                last_attempt = attempt == retries - 1

                if "429" in msg or "quota" in msg.lower() or "rate" in msg.lower():
                    wait = _retry_seconds(msg, default=15.0)
                    # A long wait means the DAILY cap is gone for this model.
                    # Retrying won't help today -> switch models immediately.
                    if wait > 45 and more_models:
                        print(f"  {current}: daily quota exhausted, "
                              f"switching to {model_queue[mi+1]}...")
                        break
                    if last_attempt:
                        if more_models:
                            print(f"  {current}: rate-limited, "
                                  f"trying {model_queue[mi+1]}...")
                            break
                        return ("[All free-tier models are rate-limited right now. "
                                "The daily quota may be used up -- it resets at "
                                "midnight US Pacific time. Try again later or add "
                                "billing in Google AI Studio.]")
                    print(f"  {current}: 429, waiting {wait:.0f}s before retry...")
                    time.sleep(wait)
                    continue

                if "503" in msg or "overloaded" in msg.lower():
                    if last_attempt and more_models:
                        break
                    time.sleep(1.5 * (attempt + 1))
                    continue

                # Unknown error: try the next model if we have one, else report.
                if more_models:
                    print(f"  {current}: error, trying next model. ({msg[:80]})")
                    break
                return f"[Could not generate a response: {msg}]"
    return "[All models exhausted. Please try again shortly.]"


# ---------------------------------------------------------------------
# Live call endpoint
# ---------------------------------------------------------------------
@app.route("/api/respond", methods=["POST"])
def respond():
    data = request.get_json(force=True)
    policy = (data.get("policy") or "").strip()
    if not policy:
        return jsonify({"error": "No policy provided."}), 400

    replies = []
    for i, slug in enumerate(AMBASSADORS):
        if i:  # small gap between representatives to respect the 5/min free tier
            time.sleep(1.2)
        text = generate(
            PRIMARY_MODEL,
            system_prompt(slug),
            f'A participant proposes this policy: "{policy}"\n'
            f"Give your constituency's position on it.",
        )
        replies.append(
            {
                "slug": slug,
                "name": EVIDENCE[slug]["display_name"],
                "data_status": EVIDENCE[slug]["data_status"],
                "text": text,
            }
        )
    return jsonify({"replies": replies})


# ---------------------------------------------------------------------
# Document analysis endpoint
# ---------------------------------------------------------------------
def extract_text(filename, raw):
    name = filename.lower()
    if name.endswith(".txt"):
        return raw.decode("utf-8", errors="ignore")
    if name.endswith(".pdf"):
        import pdfplumber
        out = []
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            for page in pdf.pages:
                out.append(page.extract_text() or "")
        return "\n".join(out)
    if name.endswith(".docx"):
        import docx
        d = docx.Document(io.BytesIO(raw))
        return "\n".join(p.text for p in d.paragraphs)
    raise ValueError("Unsupported file type. Use .txt, .docx or .pdf.")


def chunk_text(text, size=6000):
    text = text.strip()
    return [text[i : i + size] for i in range(0, len(text), size)] or [""]


def analyse_chunk(slug, chunk):
    """Return {support:[], oppose:[], improve:[]} for one ambassador/chunk."""
    prompt = f"""Analyse this section of a policy document for your constituency.

SECTION:
\"\"\"{chunk}\"\"\"

Respond ONLY with valid JSON, no markdown, in exactly this shape:
{{"support": [{{"point": "...", "confidence": "high|medium|low", "basis": "..."}}],
 "oppose":  [{{"point": "...", "confidence": "high|medium|low", "basis": "..."}}],
 "improve": [{{"point": "...", "confidence": "high|medium|low", "basis": "..."}}]}}
Use empty lists where you have nothing. 'basis' names the evidence you relied on."""
    raw = generate(PRIMARY_MODEL, system_prompt(slug), prompt)
    raw = raw.replace("```json", "").replace("```", "").strip()
    try:
        return json.loads(raw)
    except Exception:
        return {"support": [], "oppose": [], "improve": [],
                "_note": "Could not parse model output for this section."}


def build_report(filename, per_ambassador):
    import docx
    from docx.shared import Pt

    doc = docx.Document()
    doc.add_heading("Multi-Species Policy Analysis", level=0)
    doc.add_paragraph(f"Source document: {filename}")
    doc.add_paragraph(
        "Prepared by the virtual town hall. Each representative argues from its "
        "own constituency's self-interest; conflicts between them are intentional "
        "and are resolved at a system layer above these representatives."
    )

    for slug, sections in per_ambassador.items():
        a = EVIDENCE[slug]
        doc.add_heading(f"{a['display_name']}  ({a['data_status']})", level=1)

        for label, key in [
            ("Support points", "support"),
            ("Opposition points", "oppose"),
            ("Proposed improvements", "improve"),
        ]:
            doc.add_heading(label, level=2)
            rows = [pt for sec in sections for pt in sec.get(key, [])]
            if not rows:
                doc.add_paragraph("None raised.")
                continue
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = "Point", "Confidence", "Evidence basis"
            for pt in rows:
                c = table.add_row().cells
                c[0].text = str(pt.get("point", ""))
                c[1].text = str(pt.get("confidence", ""))
                c[2].text = str(pt.get("basis", ""))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@app.route("/api/analyse", methods=["POST"])
def analyse():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded."}), 400
    f = request.files["file"]
    raw = f.read()
    try:
        text = extract_text(f.filename, raw)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400

    chunks = chunk_text(text)
    per_ambassador = {}
    for slug in AMBASSADORS:
        per_ambassador[slug] = [analyse_chunk(slug, ch) for ch in chunks]

    report = build_report(f.filename, per_ambassador)
    return send_file(
        report,
        as_attachment=True,
        download_name="policy_analysis.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# ---------------------------------------------------------------------
@app.route("/")
def index():
    return send_from_directory(".", "index.html")


# Serve static assets (the tile photos: shark.jpg, sea_lion.jpg, etc.)
# from the same folder as this script.
@app.route("/<path:filename>")
def static_files(filename):
    if filename.lower().endswith((".jpg", ".jpeg", ".png", ".gif", ".webp",
                                   ".css", ".js", ".ico", ".svg")):
        return send_from_directory(".", filename)
    return ("Not found", 404)


def list_available_models():
    """Print models this key can actually use, so names are never a guess."""
    try:
        print("\nModels available to your key (that support text generation):")
        found = []
        for m in genai.list_models():
            if "generateContent" in getattr(m, "supported_generation_methods", []):
                name = m.name.replace("models/", "")
                found.append(name)
                print("   -", name)
        # Warn if our configured primary isn't in the list.
        if PRIMARY_MODEL not in found:
            print(f"\n!! Configured PRIMARY_MODEL '{PRIMARY_MODEL}' is NOT in the "
                  f"list above.\n   Edit PRIMARY_MODEL near the top of server.py to "
                  f"one of the names listed here (a 'flash-lite' one is best).")
        print()
    except Exception as e:
        print(f"(Could not list models: {e})\n")


if __name__ == "__main__":
    if API_KEY == "PASTE_YOUR_KEY_HERE":
        print("!! Set your Gemini API key in server.py or GEMINI_API_KEY first.")
    else:
        list_available_models()
    app.run(host="127.0.0.1", port=5000, debug=True)